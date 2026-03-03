"""File open/save/export actions for the canvas feature."""
from __future__ import annotations

from dataclasses import dataclass
import html
import os
from typing import TYPE_CHECKING

from PySide6.QtGui import QTextDocument, QTextListFormat
from PySide6.QtWidgets import (
    QCheckBox,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QMessageBox,
    QDoubleSpinBox,
    QSpinBox,
    QVBoxLayout,
    QWidget,
)

from services.highlights import HighlightMatch, get_highlight_store

if TYPE_CHECKING:
    from widgets.markdown.editor import EditorPanel, TabbedEditorWidget


@dataclass(slots=True)
class ExportOptions:
    """User-chosen options for document export."""

    output_format: str = "pdf"
    multi_column: bool = False
    include_highlights: bool = False
    include_comments: bool = False
    font_name: str = "Calibri"
    font_size_pt: int = 11
    line_spacing: float = 1.15


class ExportOptionsDialog(QDialog):
    """Simple options menu for PDF/Word export settings."""

    def __init__(
        self,
        parent: QWidget | None = None,
        default_format: str = "pdf",
    ):
        super().__init__(parent)
        self.setWindowTitle("Export Optionen")
        self.setModal(True)
        self.resize(460, 260)

        root = QVBoxLayout(self)
        root.setContentsMargins(12, 12, 12, 12)
        root.setSpacing(8)

        form = QFormLayout()
        form.setHorizontalSpacing(10)
        form.setVerticalSpacing(8)

        self.format_combo = QComboBox()
        self.format_combo.addItem("PDF", "pdf")
        self.format_combo.addItem("Word (DOCX)", "word")
        fmt = str(default_format or "pdf").strip().lower()
        self.format_combo.setCurrentIndex(1 if fmt == "word" else 0)
        form.addRow("Format:", self.format_combo)

        self.font_combo = QComboBox()
        self.font_combo.setEditable(True)
        for name in (
            "Calibri",
            "Arial",
            "Times New Roman",
            "Cambria",
            "Verdana",
            "Tahoma",
            "Georgia",
            "Liberation Sans",
            "Liberation Serif",
            "DejaVu Sans",
            "DejaVu Serif",
        ):
            self.font_combo.addItem(name)
        self.font_combo.setCurrentText("Calibri")
        form.addRow("Schriftart:", self.font_combo)

        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(6, 72)
        self.font_size_spin.setValue(11)
        self.font_size_spin.setSuffix(" pt")
        form.addRow("Schriftgröße:", self.font_size_spin)

        self.line_spacing_spin = QDoubleSpinBox()
        self.line_spacing_spin.setRange(1.0, 3.0)
        self.line_spacing_spin.setSingleStep(0.05)
        self.line_spacing_spin.setDecimals(2)
        self.line_spacing_spin.setValue(1.15)
        form.addRow("Zeilenabstand:", self.line_spacing_spin)

        root.addLayout(form)

        self.multi_column_cb = QCheckBox("Multi-Column Export (2 Spalten)")
        self.highlights_cb = QCheckBox("Markierungen übernehmen")
        self.comments_cb = QCheckBox(
            "Kommentare übernehmen"
        )

        root.addWidget(self.multi_column_cb)
        root.addWidget(self.highlights_cb)
        root.addWidget(self.comments_cb)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok
            | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        root.addWidget(buttons)

    def options(self) -> ExportOptions:
        fmt = str(self.format_combo.currentData() or "pdf").strip().lower()
        return ExportOptions(
            output_format="word" if fmt == "word" else "pdf",
            multi_column=self.multi_column_cb.isChecked(),
            include_highlights=self.highlights_cb.isChecked(),
            include_comments=self.comments_cb.isChecked(),
            font_name=self.font_combo.currentText().strip() or "Calibri",
            font_size_pt=self.font_size_spin.value(),
            line_spacing=float(self.line_spacing_spin.value()),
        )


class CanvasFileActions:
    """Encapsulates file dialogs and export logic for canvas tabs."""

    def __init__(self, parent: QWidget, tabs: "TabbedEditorWidget"):
        self._parent = parent
        self._tabs = tabs

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self._parent,
            "Open File",
            "",
            "Markdown (*.md *.markdown);;Text (*.txt);;All Files (*)",
        )
        if path:
            self._tabs.add_file_tab(path)

    def save_current(self):
        panel = self._current_panel()
        if panel is None:
            return

        path = panel.file_path or ""
        if not path:
            path, _ = QFileDialog.getSaveFileName(
                self._parent,
                "Save As",
                "untitled.md",
                "Markdown (*.md);;Text (*.txt);;All Files (*)",
            )
        if not path:
            return

        if panel.editor.save_file(path):
            panel.file_path = path
            idx = self._tabs.tab_widget.currentIndex()
            self._tabs.set_tab_full_title(idx, os.path.basename(path))

    def export_document(self, default_format: str = "pdf"):
        panel = self._current_panel()
        if panel is None:
            return

        options = self._ask_export_options(default_format)
        if options is None:
            return

        out_format = str(options.output_format or "pdf").strip().lower()
        is_word = out_format == "word"
        suffix = ".docx" if is_word else ".pdf"
        file_filter = "Word Document (*.docx)" if is_word else "PDF (*.pdf)"
        title = "Export as Word" if is_word else "Export as PDF"
        stem = os.path.splitext(panel.file_path or "untitled")[0]
        path, _ = QFileDialog.getSaveFileName(
            self._parent,
            title,
            stem + suffix,
            file_filter,
        )
        if not path:
            return

        try:
            idx = self._tabs.tab_widget.currentIndex()
            tab_name = self._tabs.get_tab_full_title(idx) if idx >= 0 else ""
            if is_word:
                try:
                    import docx  # noqa: F401
                except ImportError:
                    QMessageBox.warning(
                        self._parent,
                        "Missing Dependency",
                        (
                            "python-docx ist nicht installiert.\n"
                            "Bitte ausführen:  pip install python-docx"
                        ),
                    )
                    return
                self._write_docx(
                    panel.editor.toPlainText(),
                    path,
                    options=options,
                    panel_scope="draft",
                    tab_name=tab_name,
                )
            else:
                self._write_pdf(
                    panel.editor.toPlainText(),
                    path,
                    options=options,
                    panel_scope="draft",
                    tab_name=tab_name,
                )
        except Exception as exc:
            kind = "Word" if is_word else "PDF"
            QMessageBox.warning(self._parent, f"{kind} Export Failed", str(exc))

    def export_pdf(self):
        self.export_document(default_format="pdf")

    def export_word(self):
        self.export_document(default_format="word")

    def _current_panel(self) -> "EditorPanel | None":
        return self._tabs.current_panel()

    def _ask_export_options(self, default_format: str) -> ExportOptions | None:
        dialog = ExportOptionsDialog(
            self._parent,
            default_format=default_format,
        )
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return None
        return dialog.options()

    @staticmethod
    def _write_docx(
        md_text: str,
        path: str,
        *,
        options: ExportOptions,
        panel_scope: str,
        tab_name: str,
    ):
        import docx
        from docx.enum.text import WD_COLOR_INDEX
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        doc = docx.Document()
        CanvasFileActions._apply_document_typography(
            doc,
            options=options,
            qn=qn,
            OxmlElement=OxmlElement,
            Pt=Pt,
        )

        if options.multi_column:
            section = doc.sections[0]
            cols_nodes = section._sectPr.xpath("./w:cols")
            cols = cols_nodes[0] if cols_nodes else OxmlElement("w:cols")
            cols.set(qn("w:num"), "2")
            cols.set(qn("w:space"), "720")
            if not cols_nodes:
                section._sectPr.append(cols)

        parsed = CanvasFileActions._parse_markdown_lines(md_text)
        matches = CanvasFileActions._resolve_matches_for_parsed(
            parsed,
            options=options,
            panel_scope=panel_scope,
            tab_name=tab_name,
        )

        commented_ids: set[str] = set()
        char_pos = 0
        for idx, (text, style) in enumerate(parsed):
            paragraph = CanvasFileActions._new_doc_paragraph(
                doc,
                text=text,
                style=style,
                options=options,
                Pt=Pt,
            )
            para_start = char_pos
            para_end = para_start + len(text)
            local_matches = [
                item for item in matches
                if item.end > para_start and item.start < para_end
            ]

            if local_matches and (options.include_highlights or options.include_comments):
                paragraph.clear()
                CanvasFileActions._apply_paragraph_typography(
                    paragraph,
                    options=options,
                )
                for segment, active in CanvasFileActions._segments_for_block(
                    text=text,
                    para_start=para_start,
                    local_matches=local_matches,
                ):
                    if not segment:
                        continue
                    run = paragraph.add_run(segment)
                    CanvasFileActions._apply_run_typography(
                        run,
                        options=options,
                        Pt=Pt,
                    )
                    if active is None:
                        continue
                    if options.include_highlights:
                        run.font.highlight_color = (
                            CanvasFileActions._to_word_highlight_color(
                                WD_COLOR_INDEX,
                                active.color,
                            )
                        )
                    if (
                        options.include_comments
                        and active.highlight_id not in commented_ids
                    ):
                        comment_text = str(active.hover_text or "").strip()
                        if comment_text:
                            if hasattr(doc, "add_comment"):
                                doc.add_comment(
                                    run,
                                    text=comment_text,
                                    author="draft2craift",
                                    initials="D2C",
                                )
                            else:
                                doc.add_paragraph(f"Kommentar: {comment_text}")
                            commented_ids.add(active.highlight_id)

            char_pos = para_end
            if idx + 1 < len(parsed):
                char_pos += 1
        doc.save(path)

    @staticmethod
    def _write_pdf(
        md_text: str,
        path: str,
        *,
        options: ExportOptions,
        panel_scope: str,
        tab_name: str,
    ):
        from PySide6.QtPrintSupport import QPrinter

        parsed = CanvasFileActions._parse_markdown_lines(md_text)
        matches = CanvasFileActions._resolve_matches_for_parsed(
            parsed,
            options=options,
            panel_scope=panel_scope,
            tab_name=tab_name,
        )
        html_text = CanvasFileActions._build_pdf_html(
            parsed=parsed,
            matches=matches,
            options=options,
        )

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(path)

        doc = QTextDocument()
        doc.setHtml(html_text)
        doc.print_(printer)

    @staticmethod
    def _resolve_matches_for_parsed(
        parsed: list[tuple[str, str]],
        *,
        options: ExportOptions,
        panel_scope: str,
        tab_name: str,
    ) -> list[HighlightMatch]:
        if not (options.include_highlights or options.include_comments):
            return []
        plain_text = "\n".join(item[0] for item in parsed)
        return get_highlight_store().resolve_matches(
            panel_scope=panel_scope,
            tab_name=tab_name,
            full_text=plain_text,
        )

    @staticmethod
    def _segments_for_block(
        *,
        text: str,
        para_start: int,
        local_matches: list[HighlightMatch],
    ) -> list[tuple[str, HighlightMatch | None]]:
        if not local_matches:
            return [(text, None)]
        boundaries = {0, len(text)}
        for item in local_matches:
            boundaries.add(max(0, item.start - para_start))
            boundaries.add(min(len(text), item.end - para_start))
        points = sorted(boundaries)
        out: list[tuple[str, HighlightMatch | None]] = []
        for left, right in zip(points, points[1:]):
            if right <= left:
                continue
            active = CanvasFileActions._match_for_segment(
                local_matches,
                para_start + left,
                para_start + right,
            )
            out.append((text[left:right], active))
        return out

    @staticmethod
    def _build_pdf_html(
        *,
        parsed: list[tuple[str, str]],
        matches: list[HighlightMatch],
        options: ExportOptions,
    ) -> str:
        safe_font = str(options.font_name or "Calibri").replace("'", "\\'")
        line_height = max(1.0, float(options.line_spacing))
        css = (
            "<style>"
            "body { margin: 20px; color: #111111; font-family: '"
            f"{safe_font}"
            "'; font-size: "
            f"{int(options.font_size_pt)}pt;"
            " }"
            ".content { line-height: "
            f"{line_height:.2f};"
            " }"
            ".content.multi { column-count: 2; column-gap: 24pt; }"
            "h1,h2,h3,p { margin: 0 0 8pt 0; }"
            "ul,ol { margin: 0 0 8pt 18pt; padding: 0; }"
            ".comment-section { margin-top: 14pt; }"
            ".comment-list { margin: 0 0 0 18pt; padding: 0; }"
            "</style>"
        )

        out: list[str] = [
            "<html><head>",
            css,
            "</head><body>",
            (
                '<div class="content multi">'
                if options.multi_column
                else '<div class="content">'
            ),
        ]

        comment_numbers: dict[str, int] = {}
        comment_entries: list[tuple[int, str]] = []
        list_mode = ""
        char_pos = 0
        for idx, (text, style) in enumerate(parsed):
            para_start = char_pos
            para_end = para_start + len(text)
            local_matches = [
                item for item in matches
                if item.end > para_start and item.start < para_end
            ]

            content_parts: list[str] = []
            for segment, active in CanvasFileActions._segments_for_block(
                text=text,
                para_start=para_start,
                local_matches=local_matches,
            ):
                if not segment:
                    continue
                seg_html = html.escape(segment)
                if active is not None and options.include_highlights:
                    seg_html = (
                        '<span style="background-color: '
                        f"{CanvasFileActions._css_color(active.color)}"
                        ';">'
                        f"{seg_html}</span>"
                    )
                if active is not None and options.include_comments:
                    comment_text = str(active.hover_text or "").strip()
                    if comment_text and active.highlight_id not in comment_numbers:
                        number = len(comment_entries) + 1
                        comment_numbers[active.highlight_id] = number
                        comment_entries.append((number, comment_text))
                        seg_html += f'<sup>[{number}]</sup>'
                content_parts.append(seg_html)
            content = "".join(content_parts) or "&nbsp;"

            if style in {"bullet", "number"}:
                wanted_mode = "ul" if style == "bullet" else "ol"
                if list_mode != wanted_mode:
                    if list_mode:
                        out.append(f"</{list_mode}>")
                    out.append(f"<{wanted_mode}>")
                    list_mode = wanted_mode
                out.append(f"<li>{content}</li>")
            else:
                if list_mode:
                    out.append(f"</{list_mode}>")
                    list_mode = ""
                tag = "p"
                if style == "h1":
                    tag = "h1"
                elif style == "h2":
                    tag = "h2"
                elif style == "h3":
                    tag = "h3"
                out.append(f"<{tag}>{content}</{tag}>")

            char_pos = para_end
            if idx + 1 < len(parsed):
                char_pos += 1
        if list_mode:
            out.append(f"</{list_mode}>")

        out.append("</div>")
        if options.include_comments and comment_entries:
            out.append('<div class="comment-section">')
            out.append("<h3>Kommentare</h3>")
            out.append('<ol class="comment-list">')
            for number, text in comment_entries:
                out.append(
                    f"<li>[{number}] {html.escape(str(text or ''))}</li>"
                )
            out.append("</ol>")
            out.append("</div>")
        out.append("</body></html>")
        return "".join(out)

    @staticmethod
    def _css_color(value: str) -> str:
        text = str(value or "").strip().lstrip("#")
        if len(text) == 3:
            text = "".join(ch * 2 for ch in text)
        if len(text) != 6:
            return "#F9E2AF"
        try:
            int(text, 16)
        except Exception:
            return "#F9E2AF"
        return f"#{text}"

    @staticmethod
    def _new_doc_paragraph(
        doc,
        *,
        text: str,
        style: str,
        options: ExportOptions,
        Pt,
    ):
        if style == "h1":
            paragraph = doc.add_heading(text, level=1)
        elif style == "h2":
            paragraph = doc.add_heading(text, level=2)
        elif style == "h3":
            paragraph = doc.add_heading(text, level=3)
        elif style == "bullet":
            paragraph = doc.add_paragraph(text, style="List Bullet")
        elif style == "number":
            paragraph = doc.add_paragraph(text, style="List Number")
        else:
            paragraph = doc.add_paragraph(text)

        CanvasFileActions._apply_paragraph_typography(paragraph, options=options)
        for run in list(paragraph.runs):
            CanvasFileActions._apply_run_typography(
                run,
                options=options,
                Pt=Pt,
            )
        return paragraph

    @staticmethod
    def _apply_document_typography(
        doc,
        *,
        options: ExportOptions,
        qn,
        OxmlElement,
        Pt,
    ):
        for style_name in (
            "Normal",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "List Bullet",
            "List Number",
        ):
            try:
                style = doc.styles[style_name]
            except Exception:
                continue
            CanvasFileActions._apply_style_typography(
                style=style,
                options=options,
                qn=qn,
                OxmlElement=OxmlElement,
                Pt=Pt,
            )

    @staticmethod
    def _apply_style_typography(
        *,
        style,
        options: ExportOptions,
        qn,
        OxmlElement,
        Pt,
    ):
        style.font.name = options.font_name
        style.font.size = Pt(options.font_size_pt)
        style.paragraph_format.line_spacing = float(options.line_spacing)

        try:
            rpr = style._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{key}"), options.font_name)
        except Exception:
            pass

    @staticmethod
    def _apply_paragraph_typography(
        paragraph,
        *,
        options: ExportOptions,
    ):
        paragraph.paragraph_format.line_spacing = float(options.line_spacing)

    @staticmethod
    def _apply_run_typography(
        run,
        *,
        options: ExportOptions,
        Pt,
    ):
        run.font.name = options.font_name
        run.font.size = Pt(options.font_size_pt)

    @staticmethod
    def _parse_markdown_lines(md_text: str) -> list[tuple[str, str]]:
        """
        Parse markdown into display blocks aligned with Qt markdown rendering.

        This mirrors the HTML preview behavior: single line breaks inside a
        paragraph are merged into the same block; only real markdown block
        boundaries create new paragraphs.
        """
        out: list[tuple[str, str]] = []
        doc = QTextDocument()
        doc.setMarkdown(str(md_text or ""))

        bullet_styles = {
            QTextListFormat.Style.ListDisc,
            QTextListFormat.Style.ListCircle,
            QTextListFormat.Style.ListSquare,
        }
        number_styles = {
            QTextListFormat.Style.ListDecimal,
            QTextListFormat.Style.ListLowerAlpha,
            QTextListFormat.Style.ListUpperAlpha,
            QTextListFormat.Style.ListLowerRoman,
            QTextListFormat.Style.ListUpperRoman,
        }

        block = doc.begin()
        while block.isValid():
            text = str(block.text() or "").strip()
            if text:
                style = "normal"
                heading_level = int(block.blockFormat().headingLevel() or 0)
                if heading_level >= 3:
                    style = "h3"
                elif heading_level == 2:
                    style = "h2"
                elif heading_level == 1:
                    style = "h1"
                else:
                    text_list = block.textList()
                    if text_list is not None:
                        list_style = text_list.format().style()
                        if list_style in bullet_styles:
                            style = "bullet"
                        elif list_style in number_styles:
                            style = "number"
                out.append((text, style))
            block = block.next()
        return out

    @staticmethod
    def _match_for_segment(
        matches: list[HighlightMatch],
        seg_start: int,
        seg_end: int,
    ) -> HighlightMatch | None:
        selected: HighlightMatch | None = None
        best_span = -1
        for item in matches:
            if item.start >= seg_end or item.end <= seg_start:
                continue
            span = item.end - item.start
            if span > best_span:
                selected = item
                best_span = span
        return selected

    @staticmethod
    def _to_word_highlight_color(color_enum, hex_color: str):
        rgb = CanvasFileActions._hex_to_rgb(hex_color)
        palette = [
            ((255, 255, 0), color_enum.YELLOW),
            ((0, 255, 0), color_enum.BRIGHT_GREEN),
            ((0, 255, 255), color_enum.TURQUOISE),
            ((255, 0, 255), color_enum.PINK),
            ((0, 0, 255), color_enum.BLUE),
            ((255, 0, 0), color_enum.RED),
            ((128, 0, 128), color_enum.VIOLET),
            ((255, 165, 0), color_enum.DARK_YELLOW),
        ]
        best = color_enum.YELLOW
        best_dist = 10**9
        for target_rgb, word_color in palette:
            dist = (
                (rgb[0] - target_rgb[0]) ** 2
                + (rgb[1] - target_rgb[1]) ** 2
                + (rgb[2] - target_rgb[2]) ** 2
            )
            if dist < best_dist:
                best_dist = dist
                best = word_color
        return best

    @staticmethod
    def _hex_to_rgb(value: str) -> tuple[int, int, int]:
        text = str(value or "").strip().lstrip("#")
        if len(text) == 3:
            text = "".join(ch * 2 for ch in text)
        if len(text) != 6:
            return (249, 226, 175)
        try:
            return (
                int(text[0:2], 16),
                int(text[2:4], 16),
                int(text[4:6], 16),
            )
        except Exception:
            return (249, 226, 175)
