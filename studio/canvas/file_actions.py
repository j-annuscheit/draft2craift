"""File open/save/export actions for canvas tabs."""
from __future__ import annotations

from dataclasses import dataclass
import os
from typing import TYPE_CHECKING

from PySide6.QtGui import QTextDocument
from PySide6.QtPrintSupport import QPrinter
from PySide6.QtWidgets import (
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QMessageBox,
    QVBoxLayout,
    QWidget,
)

if TYPE_CHECKING:
    from studio.canvas.editor_panel import EditorPanel
    from studio.canvas.tabbed_editor_widget import TabbedEditorWidget


@dataclass(slots=True)
class ExportOptions:
    output_format: str = "pdf"


class ExportOptionsDialog(QDialog):
    def __init__(self, parent: QWidget | None = None, default_format: str = "pdf"):
        super().__init__(parent)
        self.setWindowTitle("Export Optionen")
        root = QVBoxLayout(self)
        form = QFormLayout()
        self.format_combo = QComboBox()
        self.format_combo.addItem("PDF", "pdf")
        self.format_combo.addItem("Word (DOCX)", "word")
        self.format_combo.setCurrentIndex(1 if str(default_format).lower() == "word" else 0)
        form.addRow("Format:", self.format_combo)
        root.addLayout(form)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        root.addWidget(buttons)

    def options(self) -> ExportOptions:
        fmt = str(self.format_combo.currentData() or "pdf").strip().lower()
        return ExportOptions(output_format="word" if fmt == "word" else "pdf")


def _write_pdf(markdown_text: str, path: str) -> None:
    printer = QPrinter(QPrinter.PrinterMode.HighResolution)
    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
    printer.setOutputFileName(path)
    doc = QTextDocument()
    doc.setMarkdown(str(markdown_text or ""))
    doc.print_(printer)


def _write_docx(markdown_text: str, path: str) -> None:
    import docx

    doc = docx.Document()
    for line in str(markdown_text or "").splitlines():
        doc.add_paragraph(line)
    doc.save(path)


class CanvasFileActions:
    def __init__(self, parent: QWidget, tabs: "TabbedEditorWidget"):
        self._parent = parent
        self._tabs = tabs

    def open_file(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self._parent,
            "Open File",
            "",
            "Markdown (*.md *.markdown);;Text (*.txt);;All Files (*)",
        )
        if path:
            self._tabs.add_file_tab(path)

    def save_current(self) -> None:
        panel = self._current_panel()
        if panel is None:
            return
        path = panel.file_path or ""
        if not path:
            path, _ = QFileDialog.getSaveFileName(
                self._parent,
                "Save As",
                "untitled.md",
                "Markdown (*.md);;Text (*.txt);;All Files (*)",
            )
        if not path:
            return
        if panel.editor.save_file(path):
            panel.file_path = path
            idx = self._tabs.tab_widget.currentIndex()
            self._tabs.set_tab_full_title(idx, os.path.basename(path))

    def export_document(self, default_format: str = "pdf") -> bool:
        panel = self._current_panel()
        tab_name = self._tab_name_for_panel(panel) if panel is not None else ""
        return self.export_specific_panel(panel, default_format=default_format, panel_scope="draft", tab_name=tab_name)

    def export_pdf(self) -> None:
        self.export_document(default_format="pdf")

    def export_word(self) -> None:
        self.export_document(default_format="word")

    def export_specific_panel(
        self,
        panel: QWidget | None,
        *,
        default_format: str = "pdf",
        panel_scope: str = "draft",
        tab_name: str = "",
    ) -> bool:
        del panel_scope
        if panel is None:
            return False
        editor = getattr(panel, "editor", None)
        if editor is None or not hasattr(editor, "toPlainText"):
            return False
        options = self._ask_export_options(default_format)
        if options is None:
            return False
        is_word = str(options.output_format).lower() == "word"
        suffix = ".docx" if is_word else ".pdf"
        file_filter = "Word Document (*.docx)" if is_word else "PDF (*.pdf)"
        title = "Export as Word" if is_word else "Export as PDF"
        stem_source = str(getattr(panel, "file_path", "") or "").strip() or str(tab_name or self._tab_name_for_panel(panel) or "untitled")
        stem = os.path.splitext(stem_source)[0] or "untitled"
        path, _ = QFileDialog.getSaveFileName(self._parent, title, stem + suffix, file_filter)
        if not path:
            return False
        markdown_text = str(editor.toPlainText() or "")
        try:
            if is_word:
                _write_docx(markdown_text, path)
            else:
                _write_pdf(markdown_text, path)
            return True
        except ImportError:
            QMessageBox.warning(
                self._parent,
                "Missing Dependency",
                "python-docx ist nicht installiert. Bitte ausführen: pip install python-docx",
            )
            return False
        except Exception as exc:
            kind = "Word" if is_word else "PDF"
            QMessageBox.warning(self._parent, f"{kind} Export Failed", str(exc))
            return False

    def _current_panel(self) -> "EditorPanel | None":
        return self._tabs.current_panel()

    def _tab_name_for_panel(self, panel: QWidget | None) -> str:
        if panel is None:
            return ""
        tabs = getattr(self._tabs, "tab_widget", None)
        if tabs is None:
            return ""
        for idx in range(tabs.count()):
            if tabs.widget(idx) is not panel:
                continue
            try:
                return str(self._tabs.get_tab_full_title(idx) or "").strip()
            except Exception:
                return str(tabs.tabText(idx) or "").strip()
        return ""

    def _ask_export_options(self, default_format: str) -> ExportOptions | None:
        dialog = ExportOptionsDialog(self._parent, default_format=default_format)
        if dialog.exec() != dialog.DialogCode.Accepted:
            return None
        return dialog.options()
