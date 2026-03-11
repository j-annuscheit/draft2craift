"""DOCX export writer with advanced layout options."""
from __future__ import annotations

from typing import Any

from .highlight_support import (
    local_matches_for_range,
    resolve_matches_for_parsed,
    segments_for_block,
    to_word_highlight_color,
)
from .markdown_blocks import parse_markdown_lines
from .models import ExportOptions


def write_docx(
    md_text: str,
    path: str,
    *,
    options: ExportOptions,
    panel_scope: str,
    tab_name: str,
) -> None:
    """Write a markdown document as DOCX using chosen export options."""
    import docx
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = docx.Document()
    _apply_document_typography(
        doc,
        options=options,
        qn=qn,
        oxml_element=OxmlElement,
        pt_factory=Pt,
    )

    if options.multi_column:
        section = doc.sections[0]
        cols_nodes = section._sectPr.xpath("./w:cols")
        cols = cols_nodes[0] if cols_nodes else OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "720")
        if not cols_nodes:
            section._sectPr.append(cols)

    parsed = parse_markdown_lines(md_text)
    matches = resolve_matches_for_parsed(
        parsed,
        options=options,
        panel_scope=panel_scope,
        tab_name=tab_name,
    )

    commented_ids: set[str] = set()
    char_pos = 0
    for idx, (text, style) in enumerate(parsed):
        paragraph = _new_doc_paragraph(
            doc,
            text=text,
            style=style,
            options=options,
            pt_factory=Pt,
        )
        para_start = char_pos
        para_end = para_start + len(text)
        local_matches = local_matches_for_range(matches, start=para_start, end=para_end)

        if local_matches and (options.include_highlights or options.include_comments):
            paragraph.clear()
            _apply_paragraph_typography(paragraph, options=options)
            for segment, active in segments_for_block(
                text=text,
                para_start=para_start,
                local_matches=local_matches,
            ):
                if not segment:
                    continue
                run = paragraph.add_run(segment)
                _apply_run_typography(run, options=options, pt_factory=Pt)

                if active is None:
                    continue
                if options.include_highlights:
                    run.font.highlight_color = to_word_highlight_color(
                        WD_COLOR_INDEX,
                        active.color,
                    )
                if options.include_comments and active.highlight_id not in commented_ids:
                    comment_text = str(active.hover_text or "").strip()
                    if comment_text:
                        _append_comment(doc, run, comment_text)
                        commented_ids.add(active.highlight_id)

        char_pos = para_end
        if idx + 1 < len(parsed):
            char_pos += 1

    doc.save(path)


def _append_comment(doc: Any, run: Any, comment_text: str) -> None:
    """Attach a comment if supported, else append fallback text."""
    if hasattr(doc, "add_comment"):
        doc.add_comment(
            run,
            text=comment_text,
            author="draft2craift",
            initials="D2C",
        )
        return
    doc.add_paragraph(f"Kommentar: {comment_text}")


def _new_doc_paragraph(
    doc: Any,
    *,
    text: str,
    style: str,
    options: ExportOptions,
    pt_factory: Any,
) -> Any:
    """Create a paragraph with mapped markdown style."""
    if style == "h1":
        paragraph = doc.add_heading(text, level=1)
    elif style == "h2":
        paragraph = doc.add_heading(text, level=2)
    elif style == "h3":
        paragraph = doc.add_heading(text, level=3)
    elif style == "bullet":
        paragraph = doc.add_paragraph(text, style="List Bullet")
    elif style == "number":
        paragraph = doc.add_paragraph(text, style="List Number")
    else:
        paragraph = doc.add_paragraph(text)

    _apply_paragraph_typography(paragraph, options=options)
    for run in list(paragraph.runs):
        _apply_run_typography(run, options=options, pt_factory=pt_factory)
    return paragraph


def _apply_document_typography(
    doc: Any,
    *,
    options: ExportOptions,
    qn: Any,
    oxml_element: Any,
    pt_factory: Any,
) -> None:
    """Apply base typography settings to default and heading styles."""
    for style_name in (
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Number",
    ):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        _apply_style_typography(
            style=style,
            options=options,
            qn=qn,
            oxml_element=oxml_element,
            pt_factory=pt_factory,
        )


def _apply_style_typography(
    *,
    style: Any,
    options: ExportOptions,
    qn: Any,
    oxml_element: Any,
    pt_factory: Any,
) -> None:
    """Apply font family, size and line spacing to one style."""
    style.font.name = options.font_name
    style.font.size = pt_factory(options.font_size_pt)
    style.paragraph_format.line_spacing = float(options.line_spacing)

    try:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = oxml_element("w:rFonts")
            rpr.append(rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), options.font_name)
    except AttributeError:
        return


def _apply_paragraph_typography(paragraph: Any, *, options: ExportOptions) -> None:
    paragraph.paragraph_format.line_spacing = float(options.line_spacing)


def _apply_run_typography(run: Any, *, options: ExportOptions, pt_factory: Any) -> None:
    run.font.name = options.font_name
    run.font.size = pt_factory(options.font_size_pt)
